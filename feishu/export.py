"""
Export agent markdown outputs as formatted Word documents.

Produces professional .docx files with:
- Clean typography (body 11pt, headings 15-22pt)
- Proper table formatting with alternating row colors
- Page numbering and margins
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"


def _set_cell_shading(cell, color: str) -> None:
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _style_table(table) -> None:
    """Apply professional table styling."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = FONT_EN
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            # Header row shading
            if i == 0:
                _set_cell_shading(cell, "1F4E79")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif i % 2 == 0:
                _set_cell_shading(cell, "F2F7FB")


def _add_styled_paragraph(doc, text: str, style: str = "Normal") -> None:
    """Add a paragraph with proper font settings."""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text) if text else para.add_run("")
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    return para


def markdown_to_docx(md_path: Path, output_path: Path, title: str = "") -> None:
    """Convert a markdown file to a formatted Word document.

    Args:
        md_path: Path to the markdown file.
        output_path: Desired .docx output path.
        title: Document title (extracted from first # heading if empty).
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35

    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    if not title:
        for line in lines:
            m = re.match(r"^#\s+(.+)", line)
            if m:
                title = m.group(1)
                break
        title = title or md_path.stem

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = FONT_EN
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    title_para.paragraph_format.space_after = Pt(20)

    # Add a thin line separator
    sep_para = doc.add_paragraph()
    sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep_para.add_run("─" * 50)
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    sep_para.paragraph_format.space_after = Pt(16)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip the first # heading since we used it as title
        if re.match(r"^#\s+", line) and i < 3:
            i += 1
            continue

        # H2 heading
        m = re.match(r"^##\s+(.+)", line)
        if m:
            doc.add_paragraph()
            h = doc.add_heading(level=2)
            r = h.add_run(m.group(1))
            r.font.name = FONT_EN
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            r.font.size = Pt(15)
            i += 1
            continue

        # H3 heading
        m = re.match(r"^###\s+(.+)", line)
        if m:
            doc.add_paragraph()
            h = doc.add_heading(level=3)
            r = h.add_run(m.group(1))
            r.font.name = FONT_EN
            r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            r.font.size = Pt(13)
            i += 1
            continue

        # Table
        if line.startswith("|") and "|" in line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1

            # Parse table
            rows_data = []
            for tl in table_lines:
                if re.match(r"^\|[\s\-:|]+\|$", tl):  # separator line
                    continue
                cells = [c.strip() for c in tl.split("|")[1:-1]]
                rows_data.append(cells)

            if rows_data:
                max_cols = max(len(r) for r in rows_data)
                table = doc.add_table(rows=len(rows_data), cols=max_cols)
                table.style = "Table Grid"

                for ri, row_data in enumerate(rows_data):
                    for ci, cell_text in enumerate(row_data):
                        if ci < max_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            run = p.add_run(cell_text)
                            run.font.size = Pt(10)
                            run.font.name = FONT_EN
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

                _style_table(table)
                doc.add_paragraph()  # spacing after table
            continue

        # Separator line
        if re.match(r"^---+$", line.strip()):
            sep = doc.add_paragraph()
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sep.add_run("—" * 30)
            sr.font.size = Pt(8)
            sr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            sep.paragraph_format.space_before = Pt(6)
            sep.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Blockquote or note
        if line.startswith(">"):
            note_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                note_lines.append(lines[i].lstrip("> ").strip())
                i += 1
            note_text = " ".join(note_lines)
            note_para = doc.add_paragraph()
            note_para.paragraph_format.left_indent = Cm(1.0)
            nr = note_para.add_run(f"💡 {note_text}")
            nr.font.size = Pt(10)
            nr.font.italic = True
            nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            nr.font.name = FONT_EN
            nr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            code_para = doc.add_paragraph()
            code_para.paragraph_format.left_indent = Cm(1.0)
            cr = code_para.add_run(code_text)
            cr.font.name = "Consolas"
            cr.font.size = Pt(9)
            cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Regular paragraph
        if line.strip():
            # Handle bold markers
            text = line.strip()
            p = doc.add_paragraph()
            # Simple bold handling: **text**
            parts = re.split(r"(\*\*.*?\*\*)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = FONT_EN
                r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
                r.font.size = Pt(11)

        i += 1

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("— ").font.size = Pt(8)
    # Add page number field
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_page = fp.add_run()
    run_page.font.size = Pt(8)
    run_page._r.append(fld_char_begin)
    run_page._r.append(instr_text)
    run_page._r.append(fld_char_end)
    fp.add_run(" —").font.size = Pt(8)

    doc.save(str(output_path))
